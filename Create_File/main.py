from docx import Document
from docx.shared import Cm,Pt

doc = Document()

#----- Header -----#
header = doc.sections[0].header
paragraph = header.paragraphs[0]
paragraph.add_run('VHNGROUP: Integramos Seguridad y tecnologia').bold = True

#----- Title -----#
titulo = doc.add_heading(u'Tienda VHNGROUP', level=1)
font_titulo = titulo.runs[0].font
font_titulo.size = Pt(24)

#-----Paragraphs-----#
paragraph = doc.add_paragraph(
    'Soluciones integrales en:\n'
    'Seguridad por medios electricos.\n'
    'Suministros tecnologicos.\n'
    'Redes y conectividad.\n'
    'Energia solar y UPS.\n'
    )
paragraph.add_run('www.vhngroup.com/tienda\n'
                  'www.vhngroup.com/marcas\n').bold = True

#----- Redes Sociales -----#
#IG
subtitulo_ig = doc.add_heading(u'Instagram', level=2)
doc.add_paragraph()
doc.add_picture('./images/instagram.png', Cm(1))
paragraph_ig = doc.add_paragraph()
paragraph_ig.add_run('@vhngroup').bold = True

#TK
subtitulo_tk = doc.add_heading(u'TikTok', level=2)
doc.add_paragraph()
doc.add_picture('./images/tictok.png', Cm(1))
paragraph_tk = doc.add_paragraph()
paragraph_tk.add_run('@vhngroup').bold = True

#Facebook
subtitulo_fb = doc.add_heading(u'Facebook', level=2)
doc.add_paragraph()
doc.add_picture('./images/facebook.png', Cm(1))
paragraph_fb = doc.add_paragraph()
paragraph_fb.add_run('@vhngroup').bold = True

#Youtube
subtitulo_yt = doc.add_heading(u'Youtube', level=2)
doc.add_paragraph()
doc.add_picture('./images/youtube.png', Cm(1))
paragraph_yt = doc.add_paragraph()
paragraph_yt.add_run('@vhngroup').bold = True

#--- Save the document -----#
doc.save('Documento_Final.docx')